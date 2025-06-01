from langchain_community.document_loaders.base import BaseLoader
from langchain_core.documents import Document
from docx import Document as DocxDocument
import os

class CustomDocxLoader(BaseLoader):
    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self):
        try:
            doc = DocxDocument(self.file_path)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            metadata = {"source": self.file_path}
            return [Document(page_content=full_text, metadata=metadata)]
        except Exception as e:
            raise RuntimeError(f"Error reading {self.file_path}: {e}")
        
